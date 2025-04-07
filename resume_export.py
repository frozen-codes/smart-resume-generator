"""
Resume Export Module for Resume Generator
Handles exporting resumes in different formats like TXT, HTML, DOCX, and PDF
"""

import os
import re
import datetime

# Try to import optional dependencies
try:
    import markdown
    MARKDOWN_AVAILABLE = True
except ImportError:
    MARKDOWN_AVAILABLE = False

try:
    from fpdf import FPDF
    PDF_AVAILABLE = True
except ImportError:
    PDF_AVAILABLE = False

try:
    import docx
    DOCX_AVAILABLE = True
except ImportError:
    DOCX_AVAILABLE = False

def save_as_text(resume_text, filename="resume.txt"):
    """
    Save resume as plain text file
    
    Args:
        resume_text: Resume text content
        filename: Output filename
        
    Returns:
        str: Path to the saved file
    """
    try:
        with open(filename, "w", encoding="utf-8") as f:
            f.write(resume_text)
        return os.path.abspath(filename)
    except Exception as e:
        return f"Error saving text file: {str(e)}"

def save_as_html(resume_text, filename="resume.html", dark_mode=False):
    """
    Save resume as HTML file
    
    Args:
        resume_text: Resume text content (in markdown format)
        filename: Output filename
        dark_mode: Whether to use dark mode styling
        
    Returns:
        str: Path to the saved file
    """
    if not MARKDOWN_AVAILABLE:
        plain_html = f"<pre>{resume_text}</pre>"
        
        with open(filename, "w", encoding="utf-8") as f:
            f.write(f"""<!DOCTYPE html>
<html>
<head>
    <title>Resume</title>
    <meta charset="UTF-8">
    <style>
        body {{ font-family: Arial, sans-serif; line-height: 1.6; padding: 20px; }}
        pre {{ white-space: pre-wrap; }}
    </style>
</head>
<body>
    {plain_html}
</body>
</html>""")
        return os.path.abspath(filename)
    
    try:
        # Convert markdown to HTML
        html_content = markdown.markdown(resume_text)
        
        # Determine CSS based on mode
        if dark_mode:
            css = """
        body {
            font-family: 'Segoe UI', Arial, sans-serif;
            line-height: 1.6;
            padding: 40px;
            max-width: 800px;
            margin: 0 auto;
            background-color: #121212;
            color: #e0e0e0;
        }
        h1, h2, h3 {
            color: #bb86fc;
        }
        a {
            color: #03dac6;
            text-decoration: none;
        }
        a:hover {
            text-decoration: underline;
        }
        hr {
            border: none;
            border-top: 1px solid #333;
            margin: 20px 0;
        }
        ul {
            list-style-type: square;
            color: #e0e0e0;
        }
        """
        else:
            css = """
        body {
            font-family: 'Segoe UI', Arial, sans-serif;
            line-height: 1.6;
            padding: 40px;
            max-width: 800px;
            margin: 0 auto;
            color: #333;
        }
        h1, h2, h3 {
            color: #2c3e50;
        }
        h1 {
            border-bottom: 2px solid #3498db;
            padding-bottom: 10px;
        }
        h2 {
            border-bottom: 1px solid #ddd;
            padding-bottom: 5px;
        }
        a {
            color: #3498db;
            text-decoration: none;
        }
        a:hover {
            text-decoration: underline;
        }
        ul {
            list-style-type: square;
        }
        """
        
        # Create full HTML document
        html_document = f"""<!DOCTYPE html>
<html>
<head>
    <title>Resume</title>
    <meta charset="UTF-8">
    <style>
{css}
    </style>
</head>
<body>
    {html_content}
</body>
</html>"""
        
        with open(filename, "w", encoding="utf-8") as f:
            f.write(html_document)
        
        return os.path.abspath(filename)
    
    except Exception as e:
        return f"Error saving HTML file: {str(e)}"

def save_as_pdf(resume_text, filename="resume.pdf"):
    """
    Save resume as PDF file
    
    Args:
        resume_text: Resume text content
        filename: Output filename
        
    Returns:
        str: Path to the saved file
    """
    if not PDF_AVAILABLE:
        return f"Error: PDF export requires the fpdf package. Install with 'pip install fpdf'."
    
    try:
        # Create PDF object
        pdf = FPDF()
        pdf.add_page()
        
        # Set font
        pdf.set_font("Arial", size=12)
        
        # Process and add text
        lines = resume_text.split('\n')
        for line in lines:
            # Check for heading patterns
            if line.startswith('# '):
                pdf.set_font("Arial", 'B', 16)
                pdf.cell(0, 10, line[2:], ln=True)
                pdf.set_font("Arial", size=12)
            elif line.startswith('## '):
                pdf.set_font("Arial", 'B', 14)
                pdf.cell(0, 10, line[3:], ln=True)
                pdf.set_font("Arial", size=12)
            elif line.startswith('### '):
                pdf.set_font("Arial", 'B', 12)
                pdf.cell(0, 10, line[4:], ln=True)
                pdf.set_font("Arial", size=12)
            elif line.startswith('**') and line.endswith('**'):
                # Bold text
                pdf.set_font("Arial", 'B', 12)
                pdf.cell(0, 10, line.strip('*'), ln=True)
                pdf.set_font("Arial", size=12)
            elif line.startswith('-') or line.startswith('*'):
                # Bullet point
                pdf.cell(10, 10, '•', ln=0)
                pdf.cell(0, 10, line[1:].strip(), ln=True)
            elif line.strip() == '':
                # Empty line
                pdf.ln(5)
            elif '=' in line and len(line.strip()) > 10 and all(c == '=' for c in line.strip()):
                # Separator line
                pdf.line(10, pdf.get_y(), 200, pdf.get_y())
                pdf.ln(5)
            else:
                # Regular text
                pdf.multi_cell(0, 10, line)
        
        # Save PDF
        pdf.output(filename)
        return os.path.abspath(filename)
    
    except Exception as e:
        return f"Error saving PDF file: {str(e)}"

def save_as_docx(resume_text, filename="resume.docx"):
    """
    Save resume as DOCX file
    
    Args:
        resume_text: Resume text content
        filename: Output filename
        
    Returns:
        str: Path to the saved file
    """
    if not DOCX_AVAILABLE:
        return f"Error: DOCX export requires python-docx package. Install with 'pip install python-docx'."
    
    try:
        # Create document object
        doc = docx.Document()
        
        # Process markdown and add to document
        lines = resume_text.split('\n')
        for line in lines:
            # Check for heading patterns
            if line.startswith('# '):
                doc.add_heading(line[2:], level=1)
            elif line.startswith('## '):
                doc.add_heading(line[3:], level=2)
            elif line.startswith('### '):
                doc.add_heading(line[4:], level=3)
            elif line.startswith('-') or line.startswith('*'):
                # Add bullet point
                doc.add_paragraph(line[1:].strip(), style='ListBullet')
            elif line.strip() == '':
                # Empty line - skip
                continue
            elif '=' in line and len(line.strip()) > 10 and all(c == '=' for c in line.strip()):
                # Separator - add horizontal line
                doc.add_paragraph().add_run('_' * 50)
            elif '**' in line:
                # Process bold text
                p = doc.add_paragraph()
                # Split by bold markers
                parts = re.split(r'(\*\*.*?\*\*)', line)
                for part in parts:
                    if part.startswith('**') and part.endswith('**'):
                        # Bold text
                        p.add_run(part[2:-2]).bold = True
                    else:
                        # Regular text
                        p.add_run(part)
            else:
                # Regular paragraph
                doc.add_paragraph(line)
        
        # Save document
        doc.save(filename)
        return os.path.abspath(filename)
    
    except Exception as e:
        return f"Error saving DOCX file: {str(e)}" 